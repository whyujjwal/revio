"""Resume file parsing — extract raw text from PDF and DOCX files."""

from __future__ import annotations

from pathlib import Path

from app.core.logging import get_logger

logger = get_logger(__name__)


def parse_pdf(file_path: str) -> str:
    import pdfplumber

    text_parts: list[str] = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    text = "\n\n".join(text_parts)
    logger.info("pdf parsed", file=file_path, pages=len(text_parts), chars=len(text))
    return text


def parse_docx(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    text_parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    text = "\n".join(text_parts)
    logger.info("docx parsed", file=file_path, paragraphs=len(text_parts), chars=len(text))
    return text


def parse_resume_file(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
